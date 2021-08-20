#! python3
# Ch15 - Practice Project: Custom Invitations as Word Documents
# Solution by SeungWon Bae

import docx


def add_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.font.name = run.font.name
        output_run.font.highlight_color = run.font.highlight_color
        output_run.font.size = run.font.size
        output_run.style = run.style

    # Paragraph's style and alignment data
    output_para.style = paragraph.style
    output_para.alignment = paragraph.alignment
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.widow_control = paragraph.paragraph_format.widow_control


print("Reading guest names from guests.txt...")

guests = []

with open("guests.txt", 'r') as f:
    guests = f.readlines()

print("Reading invitations_template.docx...")

doc = docx.Document("invitations_template.docx")
doc2 = docx.Document()

print("Writing invitations...")

for i, guest in enumerate(guests):
    # # Add a page break from second guest
    if i > 0:
        doc2.paragraphs[-1].add_run()
        doc2.paragraphs[-1].runs[-1].add_break(docx.enum.text.WD_BREAK.PAGE)
    
    doc.paragraphs[2].runs[0].text = guest.replace("\n", "")

    for para in doc.paragraphs:
        add_para_data(doc2, para)

doc2.save("invitations.docx")

print("DONE.")